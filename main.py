# main.py

import streamlit as st
import ffmpeg
import openai
import os
import pdfkit
from docx import Document
from pydub import AudioSegment

# OpenAI API anahtarını yapılandır
openai.api_key = 'YOUR_OPENAI_API_KEY'

# Streamlit Başlık ve Açıklama bölümü
st.title("AI-Powered Video Translation App 🎉")
st.markdown("Bu uygulama ile videolarınızı seçtiğiniz dillere çevirebilirsiniz.")

# 1. Dil Seçimi yapıldı
st.header("1. Dil Seçimi")
languages = st.multiselect(
    "Çeviri yapmak istediğiniz dilleri seçin:", 
    ["Türkçe", "İngilizce", "Fransızca", "Almanca"]
)

# 2. Video Yükleme yapıldı
st.header("2. Video Yükleme")
uploaded_file = st.file_uploader("Bir video dosyası yükleyin:", type=["mp4", "mkv", "avi"])

# İşlem Adımları açıklandı
st.subheader("İşlem Aşamaları")
st.markdown("""
1. Video ses dosyasına dönüştürülüyor (ffmpeg)
2. Ses dosyası metne çevriliyor (OpenAI Whisper)
3. Metin seçilen dillere çevriliyor (ChatGPT 4-mini)
4. Çeviriler indirilebilir hale getiriliyor
""")

# Çıktı klasörü oluşturup belirlendi
if not os.path.exists("output"):
    os.mkdir("output")

# Video -> Ses Dönüşümü Fonksiyonu kodu
def video_to_audio(input_video_path, output_audio_path):
    try:
        # ffmpeg kullanarak videoyu sese dönüştür
        ffmpeg.input(input_video_path).output(output_audio_path).run()
        st.success("Video başarıyla sese dönüştürüldü!")
    except Exception as e:
        st.error(f"Video -> Ses dönüştürme hatası: {e}")

# Ses -> Metin Çevirisi (Whisper) Fonksiyonu kodu
def transcribe_audio(audio_path):
    try:
        with open(audio_path, "rb") as audio_file:
            response = openai.Audio.transcribe("whisper-1", audio_file)
            transcription = response["text"]
            st.success("Ses dosyası başarıyla metne çevrildi!")
            return transcription
    except Exception as e:
        st.error(f"Ses -> Metin çeviri hatası: {e}")

# Metin -> Çeviri (ChatGPT) Fonksiyonu kodu
def translate_text(text, target_language):
    try:
        response = openai.Completion.create(
            engine="gpt-4-mini",
            prompt=f"Translate this text to {target_language}: {text}",
            max_tokens=1000
        )
        translation = response.choices[0].text.strip()
        return translation
    except Exception as e:
        st.error(f"Çeviri hatası: {e}")

# .docx dosyası oluşturma fonksiyonu
def save_translation_as_docx(translation, language):
    doc = Document()
    doc.add_heading(f'{language} Çevirisi', 0)
    doc.add_paragraph(translation)
    file_path = f"output/{language}_translation.docx"
    doc.save(file_path)
    return file_path

# .pdf dosyası oluşturma fonksiyonu
def save_translation_as_pdf(translation, language):
    html_content = f"<h1>{language} Çevirisi</h1><p>{translation}</p>"
    pdf_file_path = f"output/{language}_translation.pdf"
    pdfkit.from_string(html_content, pdf_file_path)
    return pdf_file_path



# Video Yüklenirse İşlemleri Başlatıyor
if uploaded_file:
    # Geçici video dosyasını kaydet
    video_path = f"output/temp_video_{uploaded_file.name}"
    with open(video_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    # Video -> Ses Dönüşümü yapıyor
    audio_path = f"output/temp_audio.wav"
    video_to_audio(video_path, audio_path)

    # Ses -> Metin Çevirisi yapıyor 
    transcription = transcribe_audio(audio_path)
    st.write("Transkript:", transcription)

     # Ses Dosyasını İndirme Seçeneği
    st.download_button(
        label="Ses Dosyasını İndir",
        data=open(audio_path, "rb"),
        file_name="audio_output.wav",
        mime="audio/wav"
    )
    
    # Metin -> Çoklu Dil Çevirisi ve İndirme Linkleri belirtiliyor
    if transcription and languages:
        translations = {}
        for lang in languages:
            translations[lang] = translate_text(transcription, lang)
            st.write(f"{lang} Çevirisi:", translations[lang])
            
            # Çeviri dosyasını oluştur ve indirilebilir hale getirdi
            translation_file_path = f"output/{lang}_translation.txt"
            with open(translation_file_path, "w") as f:
                f.write(translations[lang])
            
            # İndirme Butonu 
            st.download_button(
                label=f"{lang} Çeviriyi İndir",
                data=translations[lang],
                file_name=f"{lang}_translation.txt",
                mime="text/plain"
            )
            docx_file_path = save_translation_as_docx(translations[lang], lang)
            st.download_button(
                label=f"{lang} Çeviriyi İndir (.docx)",
                data=open(docx_file_path, "rb"),
                file_name=f"{lang}_translation.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            pdf_file_path = save_translation_as_pdf(translations[lang], lang)
            st.download_button(
                label=f"{lang} Çeviriyi İndir (.pdf)",
                data=open(pdf_file_path, "rb"),
                file_name=f"{lang}_translation.pdf",
                mime="application/pdf"
            )
# Gereksiz dosyaları temizle
if os.path.exists("output/temp_video.mp4"):
    os.remove("output/temp_video.mp4")
if os.path.exists("output/temp_audio.wav"):
    os.remove("output/temp_audio.wav")
